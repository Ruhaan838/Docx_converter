from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches,Pt
from os import path
from tkinter import messagebox

class Document_File():
    def __init__(self,list1=[],list2=[],practical_num=1,crr_page=1,code_sub="",
                 enrollment_num="",sub_name="",sub_name_small="",file_location="",list_img=[]):
        
        self.list1 = list1
        self.list2 = list2
        self.practical_num = practical_num
        self.crr_page = crr_page
        self.code_sub = code_sub
        self.enrollment_num = enrollment_num
        self.sub_name = sub_name
        self.sub_name_small = sub_name_small
        self.file_location = file_location
        self.list_img = list_img
    def docx_file(self):
        doc = Document()

        for section in doc.sections:
            header = section.header
            footer = section.footer

            header_paragraphs = header.paragraphs
            header_paragraph = header_paragraphs[0] if header_paragraphs else header.add_paragraph()

            header_run = header_paragraph.add_run()
            header_run.add_text(f'Enrollment No:{self.enrollment_num}')
            header_run.add_text(f"\t {self.sub_name}")

            header_run.bold = True
            header_run.font.size = Pt(11)

            footer_paragraphs = footer.paragraphs
            footer_paragraph = footer_paragraphs[0] if footer_paragraphs else footer.add_paragraph()

            footer_run = footer_paragraph.add_run(f"2024-2025/BE_IT_CO_DIV-I/Sem-4/{self.sub_name_small}        Computer Engg. Deptt., SCET, Surat \t Page NO:   ")

            footer_run.font.size = Pt(11)

            page_num_field = OxmlElement('w:fldSimple')
            page_num_field.set(qn('w:instr'), 'PAGE')
            page_num_run = footer_paragraph.add_run()
            page_num_run._r.append(page_num_field)
            page_num_run.font.name = 'Arial'
            page_num_run.font.size = Pt(11)

        prac_num = int(self.practical_num)

        for page_no in range(min(prac_num, len(self.list1))):
            practical_name = doc.add_heading(f"Practical {page_no + 1}", level=0)
            practical_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph('Problem Statement:')
            
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
            
            para1 = self.list2[page_no]
            para = doc.add_paragraph(para1)
            
            subtitle = doc.add_paragraph('Code:')

            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'

            if not self.list1:
                messagebox.showerror("Error","The list of code files is empty.")
                return

            with open(self.list1[page_no], 'r') as file:
                para1 = file.read()

            para = doc.add_paragraph(para1)

            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'

            doc.add_page_break()

            subtitle = doc.add_paragraph('Output:')
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'


            
            # Retrieve image for the current page_no
            current_page_img = self.list_img[-(page_no + 1)]

            img_bytes = BytesIO()
            current_page_img.save(img_bytes, format='PNG')
            img_bytes = img_bytes.getvalue()

            doc.add_picture(BytesIO(img_bytes), width=Inches(7), height=Inches(3))

            if page_no < min(prac_num, len(self.list1)) - 1:
                doc.add_page_break()

        output_docx_name = f"{self.sub_name_small}-output.docx"
        
        output_docx_path = path.join(self.file_location, output_docx_name)

        doc.save(output_docx_path)
        messagebox.showinfo("Success !",f"Your File is Export Successfully in {output_docx_path}")

